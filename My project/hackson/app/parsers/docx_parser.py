"""
Word 文档解析器
"""
import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .base import (
    BaseParser, ParsedDocument, ContentBlock, 
    DocumentMetadata, BlockType
)


class DocxParser(BaseParser):
    """Word 文档解析器"""
    
    def parse(self, file_path: Path) -> ParsedDocument:
        """解析 Word 文档"""
        doc = Document(str(file_path))
        
        metadata = self._extract_metadata(doc)
        blocks = self._parse_content(doc)
        
        # 构建原始文本
        raw_content = "\n\n".join([b.content for b in blocks])
        
        return ParsedDocument(
            metadata=metadata,
            blocks=blocks,
            raw_content=raw_content
        )
    
    def _extract_metadata(self, doc: Document) -> DocumentMetadata:
        """提取文档元数据"""
        metadata = DocumentMetadata()
        
        # 从文档属性中提取
        core_props = doc.core_properties
        if core_props.title:
            metadata.title = core_props.title
        if core_props.author:
            metadata.authors = [core_props.author]
        if core_props.keywords:
            metadata.keywords = [k.strip() for k in core_props.keywords.split(',')]
        
        return metadata
    
    def _parse_content(self, doc: Document) -> List[ContentBlock]:
        """解析文档内容"""
        blocks = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检查段落样式判断类型
            style_name = para.style.name.lower() if para.style else ""
            
            # 标题样式
            if "heading" in style_name or "title" in style_name:
                level = self._get_heading_level_from_style(style_name)
                blocks.append(ContentBlock(
                    type=BlockType.HEADING,
                    content=text,
                    should_translate=True,
                    level=level
                ))
            # 代码样式
            elif "code" in style_name:
                blocks.append(ContentBlock(
                    type=BlockType.CODE,
                    content=text,
                    should_translate=False
                ))
            # 检查是否是引用
            elif self._is_reference(text):
                blocks.append(ContentBlock(
                    type=BlockType.REFERENCE,
                    content=text,
                    should_translate=False
                ))
            # 普通段落
            else:
                # 检查是否包含公式
                has_formula = self._contains_formula(text)
                blocks.append(ContentBlock(
                    type=BlockType.TEXT,
                    content=text,
                    should_translate=True,
                    metadata={"has_formula": has_formula} if has_formula else {}
                ))
        
        # 处理表格
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                blocks.append(ContentBlock(
                    type=BlockType.TABLE,
                    content=table_text,
                    should_translate=True,
                    metadata={"is_table": True}
                ))
        
        return blocks
    
    def _get_heading_level_from_style(self, style_name: str) -> int:
        """从样式名获取标题层级"""
        match = re.search(r'heading\s*(\d+)', style_name)
        if match:
            return int(match.group(1))
        if "title" in style_name:
            return 1
        return 1
    
    def _is_reference(self, text: str) -> bool:
        """判断是否是参考文献"""
        if re.match(r'^\[?\d+[\]\.)]', text):
            return True
        return False
    
    def _contains_formula(self, text: str) -> bool:
        """检查是否包含数学公式"""
        patterns = [
            r'\$\$[\s\S]*?\$\$',
            r'\$[^\$\n]+?\$',
            r'\\begin\{equation\}',
        ]
        for pattern in patterns:
            if re.search(pattern, text):
                return True
        return False
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
    
    def get_supported_extensions(self) -> List[str]:
        return [".docx"]
